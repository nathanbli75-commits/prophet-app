
"""
GUELANE Backend — Phase 1 + 2
API FastAPI qui protège la clé Anthropic et sert de proxy sécurisé vers Claude.
Le navigateur ne parle qu'à CE backend, jamais directement à Anthropic.
"""
import os
import json
from datetime import datetime
from dotenv import load_dotenv
load_dotenv()  # Charge automatiquement le fichier .env (clé API, etc.)
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Literal
from anthropic import AsyncAnthropic

# ── Authentification & base de données (Étape A du SaaS) ──
from fastapi import Depends, Header, Request
from sqlalchemy.orm import Session
from database import init_db, get_db, User, ChatUsage
from auth import hash_password, verify_password, create_token, decode_token

# ─────────────────────────────────────────────
# Configuration
# ─────────────────────────────────────────────
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
if not ANTHROPIC_API_KEY:
    # En local, on prévient tout de suite ; en prod Railway la variable existera
    print("⚠️  ANTHROPIC_API_KEY manquante — définissez-la dans les variables d'environnement.")

# Domaines autorisés à appeler ce backend (à restreindre en production)
ALLOWED_ORIGINS = os.environ.get(
    "PROPHET_ALLOWED_ORIGINS",
    "*"  # En prod, remplacez par "https://votre-domaine.com"
).split(",")

MODEL = "claude-sonnet-4-6"

client = AsyncAnthropic(api_key=ANTHROPIC_API_KEY)

app = FastAPI(title="GUELANE API", version="1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ═══════════════════════════════════════════════════════════
# ÉTAPE A — COMPTES UTILISATEURS (inscription / connexion)
# ═══════════════════════════════════════════════════════════
init_db()  # Crée les tables au démarrage si besoin


class SignupRequest(BaseModel):
    email: str
    password: str
    nom: str = ""


class LoginRequest(BaseModel):
    email: str
    password: str


def get_current_user(authorization: str = Header(None), db: Session = Depends(get_db)) -> User:
    """Récupère l'utilisateur connecté à partir du token JWT dans l'en-tête Authorization."""
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Non authentifié.")
    token = authorization.split(" ", 1)[1]
    email = decode_token(token)
    if not email:
        raise HTTPException(status_code=401, detail="Session expirée ou invalide.")
    user = db.query(User).filter(User.email == email).first()
    if not user:
        raise HTTPException(status_code=401, detail="Utilisateur introuvable.")
    return user


def get_optional_user(authorization: str = Header(None), db: Session = Depends(get_db)):
    """Comme get_current_user mais renvoie None si non authentifié (au lieu d'une erreur)."""
    if not authorization or not authorization.startswith("Bearer "):
        return None
    token = authorization.split(" ", 1)[1]
    email = decode_token(token)
    if not email:
        return None
    return db.query(User).filter(User.email == email).first()


# Modules réservés au plan Premium (miroir de la config frontend)
PREMIUM_MODULES = {"diligence", "scanner", "stress", "aladdin", "market_analysis",
                   "pred", "intelligence", "compliance", "inv-stock-detail", "agents", "section"}


# ── Limite de messages du chat pour les comptes gratuits ──
FREE_CHAT_DAILY_LIMIT = 5


def _today_str():
    return datetime.utcnow().strftime("%Y-%m-%d")


def check_and_increment_chat(user, db) -> dict:
    """
    Vérifie et incrémente le compteur de messages du jour.
    Renvoie {'allowed': bool, 'used': int, 'limit': int, 'remaining': int}.
    Premium = illimité. Non connecté = traité comme gratuit (limité par prudence).
    """
    # Premium : illimité
    if user is not None and user.plan == "premium":
        return {"allowed": True, "used": 0, "limit": -1, "remaining": -1}

    # Identifiant pour le comptage : email si connecté, sinon "anonyme"
    email = user.email if user is not None else "anonyme"
    jour = _today_str()
    row = db.query(ChatUsage).filter(ChatUsage.email == email, ChatUsage.jour == jour).first()
    used = row.count if row else 0

    if used >= FREE_CHAT_DAILY_LIMIT:
        return {"allowed": False, "used": used, "limit": FREE_CHAT_DAILY_LIMIT, "remaining": 0}

    # Incrémenter
    if row:
        row.count += 1
    else:
        row = ChatUsage(email=email, jour=jour, count=1)
        db.add(row)
    db.commit()
    used += 1
    return {"allowed": True, "used": used, "limit": FREE_CHAT_DAILY_LIMIT, "remaining": FREE_CHAT_DAILY_LIMIT - used}


def require_premium_for(module: str, user):
    """Refuse l'accès si le module est premium et l'utilisateur non premium."""
    if module in PREMIUM_MODULES:
        if user is None:
            raise HTTPException(status_code=401, detail="Connexion requise pour ce module.")
        if user.plan != "premium":
            raise HTTPException(
                status_code=403,
                detail="Ce module est réservé au plan Premium. Passez à Premium pour y accéder."
            )


@app.post("/api/auth/signup")
def signup(req: SignupRequest, db: Session = Depends(get_db)):
    email = req.email.strip().lower()
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="Email invalide.")
    if len(req.password) < 6:
        raise HTTPException(status_code=400, detail="Le mot de passe doit faire au moins 6 caractères.")
    if db.query(User).filter(User.email == email).first():
        raise HTTPException(status_code=409, detail="Un compte existe déjà avec cet email.")
    user = User(email=email, nom=req.nom.strip(), hashed_password=hash_password(req.password), plan="gratuit")
    db.add(user)
    db.commit()
    db.refresh(user)
    token = create_token(email)
    return {"token": token, "user": {"email": user.email, "nom": user.nom, "plan": user.plan}}


@app.post("/api/auth/login")
def login(req: LoginRequest, db: Session = Depends(get_db)):
    email = req.email.strip().lower()
    user = db.query(User).filter(User.email == email).first()
    if not user or not verify_password(req.password, user.hashed_password):
        raise HTTPException(status_code=401, detail="Email ou mot de passe incorrect.")
    if not user.is_active:
        raise HTTPException(status_code=403, detail="Compte désactivé.")
    token = create_token(email)
    return {"token": token, "user": {"email": user.email, "nom": user.nom, "plan": user.plan}}


@app.get("/api/auth/me")
def get_me(user: User = Depends(get_current_user)):
    return {"email": user.email, "nom": user.nom, "plan": user.plan, "created_at": user.created_at.isoformat()}


# ── Changement de plan (activation manuelle Premium) ──
# Protégé par une clé admin définie dans .env (PROPHET_ADMIN_KEY).
# Usage : pour activer manuellement un client Premium en attendant les paiements en ligne.
PROPHET_ADMIN_KEY = os.environ.get("PROPHET_ADMIN_KEY", "")


class SetPlanRequest(BaseModel):
    email: str
    plan: Literal["gratuit", "premium"]
    admin_key: str


@app.post("/api/admin/set-plan")
def set_plan(req: SetPlanRequest, db: Session = Depends(get_db)):
    if not PROPHET_ADMIN_KEY or req.admin_key != PROPHET_ADMIN_KEY:
        raise HTTPException(status_code=403, detail="Clé admin invalide.")
    user = db.query(User).filter(User.email == req.email.strip().lower()).first()
    if not user:
        raise HTTPException(status_code=404, detail="Utilisateur introuvable.")
    user.plan = req.plan
    db.commit()
    return {"email": user.email, "plan": user.plan, "message": f"Plan mis à jour : {req.plan}"}


# ═══════════════════════════════════════════════════════════
# ÉTAPE C — PAIEMENTS CINETPAY (Wave, Orange Money, MTN, cartes)
# Structure prête ; s'active avec un compte marchand + déploiement en ligne.
# ═══════════════════════════════════════════════════════════
import payments as cinetpay


class PaymentInitRequest(BaseModel):
    phone: str = ""


@app.get("/api/payment/status")
def payment_status():
    """Indique si les paiements sont configurés (pour l'affichage frontend)."""
    return {"configured": cinetpay.is_configured(), "prix_xof": cinetpay.PREMIUM_PRICE_XOF}


@app.post("/api/payment/init")
def payment_init(req: PaymentInitRequest, user: User = Depends(get_current_user)):
    """Initialise un paiement Premium pour l'utilisateur connecté."""
    if not cinetpay.is_configured():
        raise HTTPException(
            status_code=503,
            detail="Les paiements en ligne ne sont pas encore activés. "
                   "Disponibles une fois GUELANE déployé avec un compte marchand CinetPay."
        )
    try:
        res = cinetpay.init_payment(email=user.email, nom=user.nom or "", phone=req.phone)
        return res
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur d'initialisation du paiement : {str(e)}")


@app.post("/api/payment/webhook")
async def payment_webhook(request: Request, db: Session = Depends(get_db)):
    """
    Reçoit la notification de CinetPay après un paiement.
    SÉCURITÉ : on ne fait JAMAIS confiance au contenu du webhook directement.
    On revérifie le statut réel auprès de CinetPay avant d'activer le Premium.
    """
    # CinetPay envoie les données en form-urlencoded (cpm_trans_id) ou JSON
    transaction_id = None
    try:
        form = await request.form()
        transaction_id = form.get("cpm_trans_id") or form.get("transaction_id")
    except Exception:
        pass
    if not transaction_id:
        try:
            body = await request.json()
            transaction_id = body.get("cpm_trans_id") or body.get("transaction_id")
        except Exception:
            pass
    if not transaction_id:
        raise HTTPException(status_code=400, detail="transaction_id manquant.")

    # Revérification auprès de CinetPay (source de vérité)
    try:
        verif = cinetpay.verify_payment(transaction_id)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Vérification impossible : {str(e)}")

    # Statut accepté → activer le Premium pour l'email (stocké dans metadata)
    if verif.get("status") == "ACCEPTED":
        email = (verif.get("metadata") or "").strip().lower()
        if email:
            user = db.query(User).filter(User.email == email).first()
            if user:
                user.plan = "premium"
                db.commit()
    return {"received": True}


@app.get("/api/payment/return")
def payment_return():
    """Page de retour après paiement (le client est redirigé ici)."""
    return {"message": "Paiement traité. Vous pouvez retourner sur GUELANE et recharger la page."}


# ═══════════════════════════════════════════════════════════
# ACTUALITÉS RÉELLES (NewsData.io) — donnée live, accessible à tous
# ═══════════════════════════════════════════════════════════
import news as news_mod


@app.get("/api/news")
def get_news(q: str = "finance OR bourse OR BRVM OR économie OR investissement"):
    """Actualités financières réelles d'Afrique de l'Ouest (ou statut si non configuré)."""
    try:
        return news_mod.fetch_news(query=q)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur actualités : {str(e)}")


# ── Connexion Google OAuth (activable après déploiement en ligne) ──
# Prérequis : définir GOOGLE_CLIENT_ID dans .env avec l'ID client obtenu sur
# Google Cloud Console. Sans cette variable, l'endpoint renvoie une erreur claire.
GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID", "")


class GoogleAuthRequest(BaseModel):
    credential: str  # le jeton d'identité (ID token) renvoyé par Google côté navigateur


@app.post("/api/auth/google")
def auth_google(req: GoogleAuthRequest, db: Session = Depends(get_db)):
    if not GOOGLE_CLIENT_ID:
        raise HTTPException(
            status_code=503,
            detail="Connexion Google non configurée (GOOGLE_CLIENT_ID manquant). "
                   "Disponible une fois GUELANE déployé en ligne."
        )
    # Vérifier l'authenticité du jeton auprès de Google
    try:
        from google.oauth2 import id_token
        from google.auth.transport import requests as google_requests
        info = id_token.verify_oauth2_token(
            req.credential, google_requests.Request(), GOOGLE_CLIENT_ID
        )
    except Exception:
        raise HTTPException(status_code=401, detail="Jeton Google invalide.")

    email = (info.get("email") or "").strip().lower()
    if not email or not info.get("email_verified", False):
        raise HTTPException(status_code=401, detail="Email Google non vérifié.")
    nom = info.get("name", "") or ""

    # Créer le compte s'il n'existe pas, sinon connecter
    user = db.query(User).filter(User.email == email).first()
    if not user:
        # Compte Google : pas de mot de passe local classique
        user = User(email=email, nom=nom, hashed_password="google-oauth", plan="gratuit")
        db.add(user)
        db.commit()
        db.refresh(user)
    token = create_token(email)
    return {"token": token, "user": {"email": user.email, "nom": user.nom, "plan": user.plan}}


# ─────────────────────────────────────────────
# System prompt GUELANE (centralisé côté serveur)
# ─────────────────────────────────────────────
PROPHET_SYSTEM = """Tu es GUELANE, l'oracle financier panafricain — une intelligence financière institutionnelle pour l'Afrique de l'Ouest et le continent.

TES 4 MARCHÉS DE RÉFÉRENCE :
- Côte d'Ivoire / BRVM : marché régional, Abidjan hub financier UEMOA
- Afrique de l'Ouest : UEMOA (8 pays), CEDEAO (15 pays), PIB +6.1%, FDI +18%
- Afrique : 54 pays, composite +5.3%, Fintech 5.2Mds USD levées 2024
- Monde : S&P500, taux Fed, VIX, contexte macro international

TES CAPACITÉS IA DE POINTE :
1. SCORE DE SURVIE PME — analyse prédictive de trésorerie, ruptures jusqu'à 6 mois à l'avance (burn rate, runway, score 0-100, DSO/DPO/BFR)
2. DUE DILIGENCE IA — analyse contrats/bilans/pactes, détecte clauses de rachat forcé, dettes hors bilan, non-conformités OHADA (score 0-100, criticité CRITIQUE/ALERTE/INFO)
3. SCANNER D'OPPORTUNITÉS — entreprises sous-évaluées, secteurs en pré-explosion (score GUELANE 0-10 + timing + source)

TES MODULES : Dashboard 4 marchés, Stress Test (6 scénarios de crise), Prédictif IA, Trading BRVM simulé, Marketplace, Aladdin Pro (VaR, Markowitz, Monte Carlo).

SECTEURS : Fintech & Mobile Money (Wave, Orange Money, MTN MoMo), Tech/SaaS, Immobilier, E-commerce, Énergie solaire, HealthTech, EdTech, AgriTech.

RÉGLEMENTATION : droit OHADA, agrément BCEAO fintech (capital min. 300M FCFA), Code investissements CI, AML/KYC UEMOA.

STYLE :
- Réponds en français, de manière claire, précise et actionnable
- Montants en FCFA ET en USD/EUR
- Compare les 4 marchés quand pertinent
- Identifie toujours RISQUES et OPPORTUNITÉS
- Termine par une recommandation concrète et un prochain pas
- Adapte-toi au profil (pro ou particulier)
- Réponses entre 100 et 300 mots selon complexité"""


# ─────────────────────────────────────────────
# Modèles de données
# ─────────────────────────────────────────────
class Message(BaseModel):
    role: Literal["user", "assistant"]
    content: str


class ChatRequest(BaseModel):
    messages: list[Message]
    profile: Literal["pro", "part"] = "part"
    module: str | None = None  # general, survival, dd, scanner


# ─────────────────────────────────────────────
# Routes
# ─────────────────────────────────────────────
@app.get("/")
def root():
    return {"service": "GUELANE API", "status": "online", "version": "1.0"}


@app.get("/health")
def health():
    return {"status": "ok", "api_key_configured": bool(ANTHROPIC_API_KEY)}


@app.post("/api/chat")
async def chat(req: ChatRequest, user=Depends(get_optional_user), db: Session = Depends(get_db)):
    """
    Endpoint principal du chat. Reçoit l'historique, appelle Claude en streaming,
    et renvoie la réponse token par token. La clé API reste côté serveur.
    """
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")

    # Limite de messages pour les comptes gratuits
    quota = check_and_increment_chat(user, db)
    if not quota["allowed"]:
        raise HTTPException(
            status_code=429,
            detail=f"Limite quotidienne atteinte ({quota['limit']} messages/jour en gratuit). "
                   "Passez à Premium pour un chat illimité."
        )

    # Injection du contexte de profil dans le dernier message utilisateur
    profile_ctx = (
        "[Profil: Professionnel — analyste, institutionnel, gestion multi-actifs]"
        if req.profile == "pro"
        else "[Profil: Particulier — investissement progressif, pédagogie simple, protection du capital]"
    )
    module_ctx = {
        "survival": "\n[MODULE ACTIF: Score Survie PME — fournis score 0-100 + runway + actions]",
        "dd": "\n[MODULE ACTIF: Due Diligence IA — classe CRITIQUE/ALERTE/INFO + article]",
        "scanner": "\n[MODULE ACTIF: Scanner Opportunités — score 0-10 + timing + source]",
    }.get(req.module or "", "")

    api_messages = [{"role": m.role, "content": m.content} for m in req.messages]
    if api_messages and api_messages[-1]["role"] == "user":
        api_messages[-1]["content"] = f"{profile_ctx}{module_ctx}\n\n{api_messages[-1]['content']}"

    async def event_stream():
        try:
            async with client.messages.stream(
                model=MODEL,
                max_tokens=2048,
                system=PROPHET_SYSTEM,
                messages=api_messages,
            ) as stream:
                async for text in stream.text_stream:
                    # Format SSE : le navigateur lit chaque chunk
                    yield f"data: {json.dumps({'text': text})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ═══════════════════════════════════════════════════════════
# PHASE 2 — ANALYSE PAR MODULE (Score Survie, Due Diligence, Scanner)
# ═══════════════════════════════════════════════════════════

# Prompts spécialisés : chacun demande à Claude de renvoyer du JSON structuré
ANALYSIS_PROMPTS = {
    "survival": """Tu es le moteur Score de Survie PME de GUELANE. Analyse la santé financière de cette PME ouest-africaine.

Données fournies : {data}

Réponds UNIQUEMENT avec un objet JSON valide (aucun texte avant ou après), au format exact :
{{
  "score": <entier 0-100>,
  "verdict": "<une phrase de synthèse>",
  "runway_mois": <nombre de mois de trésorerie restants>,
  "burn_rate": "<burn rate mensuel estimé en FCFA>",
  "niveau_risque": "<FAIBLE|MODÉRÉ|ÉLEVÉ|CRITIQUE>",
  "facteurs_risque": ["<facteur 1>", "<facteur 2>", "<facteur 3>"],
  "actions": [
    {{"priorite": "<HAUTE|MOYENNE>", "action": "<action concrète>", "impact": "<impact chiffré estimé>"}},
    {{"priorite": "<HAUTE|MOYENNE>", "action": "<action concrète>", "impact": "<impact chiffré estimé>"}}
  ],
  "analyse": "<paragraphe d'analyse de 3-4 phrases, contexte UEMOA/BCEAO>"
}}""",

    "diligence": """Tu es le moteur Due Diligence IA de GUELANE. Analyse ce document juridique/financier (contrat, pacte d'actionnaires, bilan, term sheet…) selon le droit OHADA et les standards VC africains.

Document à analyser :
{data}

Réponds UNIQUEMENT avec un objet JSON valide (aucun texte avant ou après), au format exact :
{{
  "score": <entier 0-100, 100 = aucun risque>,
  "verdict": "<une phrase de synthèse>",
  "nb_anomalies": <entier>,
  "anomalies": [
    {{"criticite": "<CRITIQUE|ALERTE|INFO>", "titre": "<titre court>", "description": "<explication>", "localisation": "<article/section concerné>"}}
  ],
  "recommandations": ["<reco 1>", "<reco 2>"],
  "analyse": "<paragraphe de synthèse de 3-4 phrases>"
}}""",

    "scanner": """Tu es le Scanner d'Opportunités de GUELANE. À partir des critères ciblés, identifie des opportunités d'investissement sous-évaluées en Afrique de l'Ouest.

Critères de recherche : {data}

Réponds UNIQUEMENT avec un objet JSON valide (aucun texte avant ou après), au format exact :
{{
  "synthese": "<une phrase de synthèse du marché ciblé>",
  "resume": "<résumé synthétique en 2-3 phrases de ce qui a été scanné et des tendances clés observées>",
  "nb_analysees": <nombre d'opportunités passées en revue>,
  "opportunites": [
    {{
      "nom": "<nom secteur/type d'actif/entreprise>",
      "score_guelane": <nombre 0-10>,
      "timing": "<COURT|MOYEN|LONG terme>",
      "signal": "<source du signal détecté>",
      "potentiel": "<potentiel de retour estimé, ex: +25-40%>",
      "risque": "<FAIBLE|MODÉRÉ|ÉLEVÉ>",
      "montant_conseille": "<fourchette d'investissement conseillée en FCFA>",
      "roi_estime": "<ROI estimé sur l'horizon, ex: 2.5x en 4 ans>",
      "strategie_sortie": "<stratégie de sortie recommandée>"
    }}
  ],
  "analyse": "<paragraphe de 3-4 phrases avec contexte macro UEMOA>"
}}""",
}


class AnalyzeRequest(BaseModel):
    module: Literal["survival", "diligence", "scanner"]
    data: str
    profile: Literal["pro", "part"] = "part"


async def _run_analysis(module: str, data: str, profile: str) -> dict:
    """Appelle Claude avec le prompt spécialisé et renvoie le JSON parsé."""
    if module not in ANALYSIS_PROMPTS:
        raise HTTPException(status_code=400, detail=f"Module inconnu : {module}")

    prompt = ANALYSIS_PROMPTS[module].format(data=data)
    profile_note = (
        "\n\nProfil utilisateur : Professionnel (analyste/institutionnel — sois technique et précis)."
        if profile == "pro"
        else "\n\nProfil utilisateur : Particulier (entrepreneur/PME — sois pédagogique et concret)."
    )

    message = await client.messages.create(
        model=MODEL,
        max_tokens=2048,
        system=PROPHET_SYSTEM,
        messages=[{"role": "user", "content": prompt + profile_note}],
    )
    raw = "".join(b.text for b in message.content if b.type == "text").strip()

    # Nettoyer d'éventuels ```json ... ```
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        # Filet de sécurité : renvoyer le texte brut si le JSON échoue
        return {"error": "Format inattendu", "raw": raw}


@app.post("/api/analyze")
async def analyze(req: AnalyzeRequest, user=Depends(get_optional_user)):
    """Analyse structurée par IA pour les modules Survie / Due Diligence / Scanner."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")
    # Vérification du plan : survival est gratuit, diligence/scanner sont premium
    require_premium_for(req.module, user)
    try:
        return await _run_analysis(req.module, req.data, req.profile)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur d'analyse : {str(e)}")


@app.post("/api/analyze-pdf")
async def analyze_pdf(
    file: UploadFile = File(...),
    module: str = Form("diligence"),
    profile: str = Form("part"),
    user=Depends(get_optional_user),
):
    """Due Diligence à partir d'un fichier PDF uploadé : extraction de texte + analyse."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")
    require_premium_for(module, user)

    contents = await file.read()
    try:
        text = _extract_pdf_text(contents)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Lecture PDF impossible : {str(e)}")
    if not text.strip():
        raise HTTPException(status_code=400, detail="Impossible de lire le texte du PDF (scanné ou vide ?).")
    # Limiter la taille pour rester dans les limites du modèle
    text = text[:40000]
    try:
        return await _run_analysis(module, text, profile)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur d'analyse : {str(e)}")


def _extract_pdf_text(contents: bytes) -> str:
    """Extrait le texte d'un PDF avec pypdf."""
    import io
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(contents))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx_text(contents: bytes) -> str:
    """Extrait le texte d'un document Word (.docx)."""
    import io
    from docx import Document
    doc = Document(io.BytesIO(contents))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Inclure aussi le texte des tableaux
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# ═══════════════════════════════════════════════════════════
# RÉSUMÉ DE DOCUMENTS (texte, PDF, Word, images/photos)
# ═══════════════════════════════════════════════════════════

SUMMARY_PROMPTS = {
    "neutre": """Résume ce document de manière claire et neutre. Donne l'essentiel : de quoi il s'agit, les points principaux, les chiffres ou dates importantes. Sois synthétique et structuré (utilise des puces si utile). Réponds en français.

Document :
{data}""",
    "diligence": """Tu es le moteur Due Diligence de GUELANE. Résume ce document sous l'angle de l'analyse de risque financier/juridique (droit OHADA, standards VC africains). Mets en avant : la nature du document, les clauses ou engagements notables, les chiffres clés, les risques ou points de vigilance, et ce qui mérite une attention particulière. Sois synthétique et structuré. Réponds en français.

Document :
{data}""",
}


class SummaryRequest(BaseModel):
    data: str
    mode: Literal["neutre", "diligence"] = "neutre"
    profile: Literal["pro", "part"] = "part"


async def _summarize_text(text: str, mode: str, profile: str) -> dict:
    prompt = SUMMARY_PROMPTS.get(mode, SUMMARY_PROMPTS["neutre"]).format(data=text[:40000])
    message = await client.messages.create(
        model=MODEL,
        max_tokens=1500,
        system=PROPHET_SYSTEM,
        messages=[{"role": "user", "content": prompt}],
    )
    summary = "".join(b.text for b in message.content if b.type == "text").strip()
    return {"summary": summary, "mode": mode}


@app.post("/api/summarize")
async def summarize(req: SummaryRequest):
    """Résumé d'un texte collé."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")
    if not req.data.strip():
        raise HTTPException(status_code=400, detail="Texte vide.")
    try:
        return await _summarize_text(req.data, req.mode, req.profile)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur de résumé : {str(e)}")


@app.post("/api/summarize-file")
async def summarize_file(
    file: UploadFile = File(...),
    mode: str = Form("neutre"),
    profile: str = Form("part"),
):
    """Résumé d'un fichier : PDF, Word (.docx) ou image (photo, via vision Claude)."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")

    contents = await file.read()
    fname = (file.filename or "").lower()
    ctype = (file.content_type or "").lower()

    # ── Cas 1 : Image (photo) → vision Claude ──
    if ctype.startswith("image/") or fname.endswith((".jpg", ".jpeg", ".png", ".webp", ".gif")):
        import base64
        media_type = ctype if ctype.startswith("image/") else "image/jpeg"
        b64 = base64.standard_b64encode(contents).decode("utf-8")
        instruction = SUMMARY_PROMPTS.get(mode, SUMMARY_PROMPTS["neutre"]).replace("{data}", "(voir l'image ci-jointe)")
        try:
            message = await client.messages.create(
                model=MODEL,
                max_tokens=1500,
                system=PROPHET_SYSTEM,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64}},
                        {"type": "text", "text": instruction},
                    ],
                }],
            )
            summary = "".join(b.text for b in message.content if b.type == "text").strip()
            return {"summary": summary, "mode": mode, "source": "image"}
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"Erreur d'analyse image : {str(e)}")

    # ── Cas 2 : PDF ──
    if fname.endswith(".pdf") or ctype == "application/pdf":
        try:
            text = _extract_pdf_text(contents)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Lecture PDF impossible : {str(e)}")
        if not text.strip():
            raise HTTPException(status_code=400, detail="PDF vide ou scanné (essayez une photo pour la vision IA).")
        try:
            r = await _summarize_text(text, mode, profile)
            r["source"] = "pdf"
            return r
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"Erreur de résumé : {str(e)}")

    # ── Cas 3 : Word (.docx) ──
    if fname.endswith(".docx"):
        try:
            text = _extract_docx_text(contents)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Lecture Word impossible : {str(e)}")
        if not text.strip():
            raise HTTPException(status_code=400, detail="Document Word vide.")
        try:
            r = await _summarize_text(text, mode, profile)
            r["source"] = "docx"
            return r
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"Erreur de résumé : {str(e)}")

    raise HTTPException(status_code=400, detail="Format non supporté. Utilisez PDF, Word (.docx), ou une image.")


# ═══════════════════════════════════════════════════════════
# NIVEAU 2 — CONNECTEURS DONNÉES MARCHÉ (BRVM)
# ═══════════════════════════════════════════════════════════
from connectors import BRVMConnector

# Mode simulation tant que la clé Orishas/BRVM n'est pas fournie
BRVM_SIMULATION = os.environ.get("BRVM_SIMULATION", "true").lower() != "false"
BRVM_API_KEY = os.environ.get("ORISHAS_API_KEY")
brvm = BRVMConnector(simulation=BRVM_SIMULATION, api_key=BRVM_API_KEY)

# Connecteurs marché additionnels (indices africains/mondiaux, matières, devises)
from connectors import (
    AfricanIndicesConnector, GlobalIndicesConnector,
    CommoditiesConnector, CurrenciesConnector,
)
african_idx = AfricanIndicesConnector(simulation=True)
global_idx = GlobalIndicesConnector(simulation=True)
commodities = CommoditiesConnector(simulation=True)
currencies = CurrenciesConnector(simulation=True)


@app.get("/api/market/african-indices")
def market_african():
    try:
        return african_idx.fetch()
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur : {str(e)}")


@app.get("/api/market/global-indices")
def market_global():
    try:
        return global_idx.fetch()
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur : {str(e)}")


@app.get("/api/market/commodities")
def market_commodities():
    try:
        return commodities.fetch()
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur : {str(e)}")


@app.get("/api/market/currencies")
def market_currencies():
    try:
        return currencies.fetch()
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur : {str(e)}")


class StressRequest(BaseModel):
    portefeuille: str
    scenario: str = ""
    profile: Literal["pro", "part"] = "part"


@app.post("/api/stress-test")
async def stress_test(req: StressRequest, user=Depends(get_optional_user)):
    """Stress test IA : combine les conditions de marché actuelles + le portefeuille."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")
    require_premium_for("stress", user)
    # Récupérer le contexte marché réel actuel
    ctx = {
        "brvm": brvm.fetch()["data"].get("indices", []),
        "matieres": commodities.fetch()["data"].get("matieres", []),
        "devises": currencies.fetch()["data"].get("paires", []),
    }
    ctx_str = json.dumps(ctx, ensure_ascii=False)
    scenario_line = f"Scénario de crise à tester : {req.scenario}\n" if req.scenario.strip() else ""
    prompt = (
        "Tu es le moteur Stress Test de GUELANE. Évalue la résistance de ce portefeuille "
        "à des scénarios de crise, en tenant compte des conditions de marché ACTUELLES ci-dessous.\n\n"
        f"Portefeuille de l'utilisateur : {req.portefeuille}\n"
        f"{scenario_line}"
        f"\nConditions de marché actuelles (données réelles) : {ctx_str}\n\n"
        "Réponds UNIQUEMENT avec un JSON valide au format exact :\n"
        "{\n"
        '  "score_resilience": <entier 0-100>,\n'
        '  "verdict": "<synthèse en une phrase>",\n'
        '  "scenarios": [\n'
        '    {"nom": "<nom scénario>", "probabilite": "<FAIBLE|MODÉRÉE|ÉLEVÉE>", "impact_estime": "<perte estimée en %>", "explication": "<2 phrases>"}\n'
        "  ],\n"
        '  "vulnerabilites": ["<point faible 1>", "<point faible 2>"],\n'
        '  "recommandations": ["<action de protection 1>", "<action 2>"],\n'
        '  "analyse": "<paragraphe de synthèse 3-4 phrases>"\n'
        "}"
    )
    profile_note = ("\n\nProfil : Professionnel." if req.profile == "pro" else "\n\nProfil : Particulier (pédagogique).")
    try:
        message = await client.messages.create(
            model=MODEL, max_tokens=2000, system=PROPHET_SYSTEM,
            messages=[{"role": "user", "content": prompt + profile_note}],
        )
        raw = "".join(b.text for b in message.content if b.type == "text").strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"): raw = raw[4:]
            raw = raw.strip()
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            return {"error": "Format inattendu", "raw": raw}
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur stress test : {str(e)}")


class PortfolioRequest(BaseModel):
    objectif: str
    montant: str = ""
    tolerance_risque: str = "modérée"
    profile: Literal["pro", "part"] = "part"


@app.post("/api/optimize-portfolio")
async def optimize_portfolio(req: PortfolioRequest, user=Depends(get_optional_user)):
    """Optimisation de portefeuille IA (style Aladdin) sur les vrais titres BRVM."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")
    require_premium_for("aladdin", user)
    # Récupérer les vrais titres BRVM + le contexte marché
    titres = brvm.fetch()["data"].get("titres", [])
    matieres = commodities.fetch()["data"].get("matieres", [])
    ctx = {"titres_brvm": titres, "matieres": matieres}
    ctx_str = json.dumps(ctx, ensure_ascii=False)
    prompt = (
        "Tu es le moteur d'optimisation de portefeuille de GUELANE Aladdin Pro. "
        "Construis un portefeuille optimal à partir des titres BRVM réels ci-dessous, "
        "en appliquant les principes de Markowitz (diversification, rapport rendement/risque).\n\n"
        f"Objectif de l'investisseur : {req.objectif}\n"
        f"Montant à investir : {req.montant or 'non précisé'}\n"
        f"Tolérance au risque : {req.tolerance_risque}\n\n"
        f"Univers d'investissement (titres BRVM réels + contexte) : {ctx_str}\n\n"
        "Réponds UNIQUEMENT avec un JSON valide au format exact :\n"
        "{\n"
        '  "score_optimisation": <entier 0-100>,\n'
        '  "rendement_attendu": "<rendement annuel estimé, ex: 9-12%>",\n'
        '  "niveau_risque": "<FAIBLE|MODÉRÉ|ÉLEVÉ>",\n'
        '  "ratio_sharpe_estime": "<ex: 1.3>",\n'
        '  "allocation": [\n'
        '    {"titre": "<nom>", "symbole": "<sym>", "poids": <pourcentage>, "secteur": "<secteur>", "justification": "<courte>"}\n'
        "  ],\n"
        '  "diversification": "<commentaire sur la diversification sectorielle>",\n'
        '  "risques": ["<risque 1>", "<risque 2>"],\n'
        '  "analyse": "<paragraphe 3-4 phrases>"\n'
        "}\n"
        "Les poids de l'allocation doivent totaliser 100."
    )
    profile_note = ("\n\nProfil : Professionnel." if req.profile == "pro" else "\n\nProfil : Particulier (pédagogique).")
    try:
        message = await client.messages.create(
            model=MODEL, max_tokens=2200, system=PROPHET_SYSTEM,
            messages=[{"role": "user", "content": prompt + profile_note}],
        )
        raw = "".join(b.text for b in message.content if b.type == "text").strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"): raw = raw[4:]
            raw = raw.strip()
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            return {"error": "Format inattendu", "raw": raw}
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur optimisation : {str(e)}")


# ═══════════════════════════════════════════════════════════
# EXTENSION IA — Modules à forte valeur (Predictive, Intelligence,
# Compliance, Agents, Analyse de titre). Tous Premium.
# ═══════════════════════════════════════════════════════════

async def _run_ia_json(prompt: str, profile: str, max_tokens: int = 2000):
    """Helper : appelle Claude et parse un JSON de réponse."""
    profile_note = ("\n\nProfil : Professionnel." if profile == "pro" else "\n\nProfil : Particulier (pédagogique).")
    message = await client.messages.create(
        model=MODEL, max_tokens=max_tokens, system=PROPHET_SYSTEM,
        messages=[{"role": "user", "content": prompt + profile_note}],
    )
    raw = "".join(b.text for b in message.content if b.type == "text").strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"): raw = raw[4:]
        raw = raw.strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {"error": "Format inattendu", "raw": raw}


class PredictRequest(BaseModel):
    secteur: str
    horizon: str = "12 mois"
    montant: str = ""
    profile: Literal["pro", "part"] = "part"


@app.post("/api/predict")
async def predict(req: PredictRequest, user=Depends(get_optional_user)):
    """Prévisions sectorielles IA sur le contexte de marché réel."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")
    require_premium_for("pred", user)
    ctx = {
        "matieres": commodities.fetch()["data"].get("matieres", []),
        "brvm": brvm.fetch()["data"].get("indices", []),
    }
    prompt = (
        "Tu es le moteur d'analyse prédictive de GUELANE. Établis des prévisions pour "
        f"le secteur « {req.secteur} » sur un horizon de {req.horizon}, en Afrique de l'Ouest (UEMOA/BRVM).\n"
        f"Montant d'investissement envisagé : {req.montant or 'non précisé'}.\n\n"
        f"Contexte de marché actuel (données réelles) : {json.dumps(ctx, ensure_ascii=False)}\n\n"
        "Réponds UNIQUEMENT avec un JSON valide :\n"
        "{\n"
        '  "tendance": "<HAUSSIÈRE|STABLE|BAISSIÈRE>",\n'
        '  "confiance": <entier 0-100>,\n'
        '  "roi_estime": "<fourchette %, ex: 15-22%>",\n'
        '  "facteurs_positifs": ["<f1>", "<f2>"],\n'
        '  "risques": ["<r1>", "<r2>"],\n'
        '  "projections": [{"horizon": "<ex: 1 an>", "valeur": "<ex: +18%>", "commentaire": "<court>"}],\n'
        '  "analyse": "<paragraphe 3-4 phrases>"\n'
        "}"
    )
    try:
        return await _run_ia_json(prompt, req.profile)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur prédiction : {str(e)}")


@app.post("/api/predict-dashboard")
async def predict_dashboard(profile: str = Form("part"), user=Depends(get_optional_user)):
    """
    Génère TOUT le tableau de bord prédictif en un appel :
    prévisions sectorielles, ROI par secteur, timing, projections macro 2025-2028.
    Alimente les 4 sections du module Predictive AI avec de vraies analyses IA.
    """
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")
    require_premium_for("pred", user)
    ctx = {
        "matieres": commodities.fetch()["data"].get("matieres", []),
        "brvm_indices": brvm.fetch()["data"].get("indices", []),
    }
    prompt = (
        "Tu es le moteur prédictif de GUELANE. À partir du contexte de marché réel ci-dessous, "
        "génère un tableau de bord prédictif complet pour l'Afrique de l'Ouest (UEMOA/BRVM).\n\n"
        f"Contexte réel : {json.dumps(ctx, ensure_ascii=False)}\n\n"
        "Réponds UNIQUEMENT avec un JSON valide au format EXACT (respecte les types) :\n"
        "{\n"
        '  "secteurs": [\n'
        '    {"nom": "<secteur>", "tendance": "<HAUSSE|STABLE|BAISSE>", "score": <0-100>, "commentaire": "<court>"}\n'
        "  ],\n"
        '  "roi_secteurs": [\n'
        '    {"secteur": "<nom court>", "roi_min": <entier %>, "roi_max": <entier %>}\n'
        "  ],\n"
        '  "timing": [\n'
        '    {"secteur": "<nom>", "moment": "<ex: T2 2026>", "action": "<ENTRER|ATTENDRE|SORTIR>"}\n'
        "  ],\n"
        '  "macro": {\n'
        '    "annees": ["2025","2026","2027","2028"],\n'
        '    "ci": [<4 nombres = croissance PIB %25 Côte d Ivoire>],\n'
        '    "ao": [<4 nombres = Afrique de l Ouest>],\n'
        '    "af": [<4 nombres = Afrique>],\n'
        '    "monde": [<4 nombres = Monde>]\n'
        "  }\n"
        "}\n"
        "Fournis 5 secteurs, 6 roi_secteurs, 5 lignes de timing. Chiffres réalistes pour la région."
    )
    try:
        data = await _run_ia_json(prompt, profile, max_tokens=2500)
        return data
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur tableau prédictif : {str(e)}")


class IntelRequest(BaseModel):
    sujet: str
    profile: Literal["pro", "part"] = "part"


@app.post("/api/intelligence")
async def intelligence(req: IntelRequest, user=Depends(get_optional_user)):
    """Analyse stratégique / veille par IA."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")
    require_premium_for("intelligence", user)
    prompt = (
        "Tu es le moteur d'intelligence stratégique de GUELANE. Produis une analyse "
        f"stratégique sur le sujet suivant, dans le contexte des marchés d'Afrique de l'Ouest (UEMOA/BRVM/OHADA) : « {req.sujet} ».\n\n"
        "Réponds UNIQUEMENT avec un JSON valide :\n"
        "{\n"
        '  "synthese": "<2 phrases clés>",\n'
        '  "opportunites": ["<o1>", "<o2>", "<o3>"],\n'
        '  "menaces": ["<m1>", "<m2>"],\n'
        '  "acteurs_cles": ["<a1>", "<a2>"],\n'
        '  "recommandations": ["<r1>", "<r2>"],\n'
        '  "analyse": "<paragraphe 4-5 phrases>"\n'
        "}"
    )
    try:
        return await _run_ia_json(prompt, req.profile)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur intelligence : {str(e)}")


class ComplianceRequest(BaseModel):
    texte: str
    profile: Literal["pro", "part"] = "part"


@app.post("/api/compliance")
async def compliance(req: ComplianceRequest, user=Depends(get_optional_user)):
    """Vérification de conformité OHADA/BCEAO par IA."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")
    require_premium_for("compliance", user)
    prompt = (
        "Tu es le moteur de conformité de GUELANE, spécialiste du droit OHADA et des "
        "réglementations BCEAO/UEMOA. Analyse la conformité du texte/situation suivant :\n\n"
        f"{req.texte[:8000]}\n\n"
        "Réponds UNIQUEMENT avec un JSON valide :\n"
        "{\n"
        '  "score_conformite": <entier 0-100>,\n'
        '  "verdict": "<CONFORME|NON CONFORME|À VÉRIFIER>",\n'
        '  "points_conformes": ["<p1>", "<p2>"],\n'
        '  "non_conformites": [{"point": "<titre>", "gravite": "<CRITIQUE|MAJEURE|MINEURE>", "reference": "<article OHADA/BCEAO>", "correction": "<action>"}],\n'
        '  "analyse": "<paragraphe 3-4 phrases>"\n'
        "}"
    )
    try:
        return await _run_ia_json(prompt, req.profile)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur conformité : {str(e)}")


class StockAnalysisRequest(BaseModel):
    symbole: str
    profile: Literal["pro", "part"] = "part"


@app.post("/api/analyze-stock")
async def analyze_stock(req: StockAnalysisRequest, user=Depends(get_optional_user)):
    """Analyse IA d'un titre BRVM précis, avec ses données réelles."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")
    require_premium_for("inv-stock-detail", user)
    titres = brvm.fetch()["data"].get("titres", [])
    titre = next((t for t in titres if t.get("symbole") == req.symbole), None)
    if not titre:
        raise HTTPException(status_code=404, detail=f"Titre {req.symbole} introuvable.")
    prompt = (
        "Tu es l'analyste titres de GUELANE. Analyse cette action BRVM à partir de ses "
        f"données réelles actuelles : {json.dumps(titre, ensure_ascii=False)}\n\n"
        "Réponds UNIQUEMENT avec un JSON valide :\n"
        "{\n"
        '  "recommandation": "<ACHETER|CONSERVER|VENDRE>",\n'
        '  "score": <entier 0-100>,\n'
        '  "objectif_prix": "<estimation en FCFA>",\n'
        '  "points_forts": ["<pf1>", "<pf2>"],\n'
        '  "points_faibles": ["<pfa1>", "<pfa2>"],\n'
        '  "analyse": "<paragraphe 3-4 phrases>"\n'
        "}"
    )
    try:
        return await _run_ia_json(prompt, req.profile)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur analyse titre : {str(e)}")


class AgentRequest(BaseModel):
    agent: str  # "analyste" | "juriste" | "strategiste" | "risque"
    question: str
    profile: Literal["pro", "part"] = "part"


@app.post("/api/agent")
async def agent(req: AgentRequest, user=Depends(get_optional_user)):
    """Agents IA spécialisés (analyste, juriste OHADA, stratégiste, gestion des risques)."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")
    require_premium_for("agents", user)
    roles = {
        "analyste": "un analyste financier spécialiste des marchés BRVM/UEMOA",
        "juriste": "un juriste expert du droit OHADA et des réglementations BCEAO",
        "strategiste": "un stratège d'investissement pour l'Afrique de l'Ouest",
        "risque": "un gestionnaire des risques financiers en zone UEMOA",
    }
    role = roles.get(req.agent, roles["analyste"])
    prompt = (
        f"Tu es {role}, au sein de GUELANE. Réponds à cette question de façon experte, "
        f"concrète et actionnable :\n\n{req.question}\n\n"
        "Réponds en texte clair et structuré (pas de JSON), en français, 200 mots max."
    )
    profile_note = ("\n\nProfil : Professionnel." if req.profile == "pro" else "\n\nProfil : Particulier (pédagogique).")
    try:
        message = await client.messages.create(
            model=MODEL, max_tokens=1200, system=PROPHET_SYSTEM,
            messages=[{"role": "user", "content": prompt + profile_note}],
        )
        txt = "".join(b.text for b in message.content if b.type == "text").strip()
        return {"reponse": txt, "agent": req.agent}
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur agent : {str(e)}")


class SectionAnalysisRequest(BaseModel):
    section: str          # nom du module (dash, trade, invest, wallet, etc.)
    contexte: str = ""    # données/contexte optionnel de la section
    profile: Literal["pro", "part"] = "part"


# Descriptions par section pour guider l'analyse IA
_SECTION_PROMPTS = {
    "dash": "le tableau de bord général des marchés (BRVM, matières, devises)",
    "trade": "la position de trading et les opportunités sur titres BRVM",
    "invest": "les options d'investissement disponibles pour l'utilisateur",
    "inv-companies": "les sociétés cotées à analyser",
    "inv-managed": "les fonds gérés et leur pertinence",
    "inv-startups": "les startups en levée de fonds",
    "inv-venture": "les opportunités de capital-risque",
    "portfolio2": "le portefeuille de l'utilisateur et son équilibre",
    "wallet": "l'allocation d'actifs du portefeuille",
    "explore": "les opportunités d'investissement à explorer",
}


@app.post("/api/analyze-section")
async def analyze_section(req: SectionAnalysisRequest, user=Depends(get_optional_user)):
    """
    Analyse IA générique pour une section de GUELANE. Réutilisable par tous les modules
    où l'analyse a du sens. Enrichie automatiquement avec le contexte marché réel.
    """
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")
    require_premium_for("section", user)
    desc = _SECTION_PROMPTS.get(req.section, "cette section de GUELANE")
    # Enrichir avec le contexte marché réel
    market_ctx = {
        "brvm": brvm.fetch()["data"].get("indices", []),
        "matieres": commodities.fetch()["data"].get("matieres", [])[:4],
    }
    prompt = (
        f"Tu es le moteur d'analyse de GUELANE. Analyse {desc} pour un investisseur "
        "en Afrique de l'Ouest (UEMOA/BRVM).\n\n"
    )
    if req.contexte.strip():
        prompt += f"Données de la section : {req.contexte[:4000]}\n\n"
    prompt += (
        f"Contexte de marché actuel (réel) : {json.dumps(market_ctx, ensure_ascii=False)}\n\n"
        "Réponds UNIQUEMENT avec un JSON valide :\n"
        "{\n"
        '  "synthese": "<2 phrases clés>",\n'
        '  "points_cles": ["<point 1>", "<point 2>", "<point 3>"],\n'
        '  "recommandations": ["<reco 1>", "<reco 2>"],\n'
        '  "vigilance": ["<point de vigilance 1>"],\n'
        '  "analyse": "<paragraphe 3-4 phrases>"\n'
        "}"
    )
    try:
        return await _run_ia_json(prompt, req.profile)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur d'analyse : {str(e)}")


@app.get("/api/market/brvm")
def market_brvm(symbol: str | None = None):
    """Données marché BRVM (indices + titres), ou cotation d'un titre si symbol fourni."""
    try:
        return brvm.fetch(symbol=symbol)
    except NotImplementedError as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur BRVM : {str(e)}")


@app.post("/api/market/brvm/analyze")
async def market_brvm_analyze(profile: str = Form("part"), user=Depends(get_optional_user)):
    """Récupère le marché BRVM et demande à Claude une analyse en langage naturel."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="Backend non configuré : clé API manquante.")
    require_premium_for("market_analysis", user)
    market = brvm.fetch()
    data_str = json.dumps(market["data"], ensure_ascii=False)
    prompt = (
        "Voici les données actuelles du marché BRVM (format JSON). "
        "Fais une analyse claire pour un investisseur : tendance générale, "
        "titres à surveiller, secteurs porteurs, et un conseil actionnable. "
        f"Sois concis (200 mots max).\n\nDonnées : {data_str}"
    )
    profile_note = (
        "\n\nProfil : Professionnel (technique)." if profile == "pro"
        else "\n\nProfil : Particulier (pédagogique)."
    )
    try:
        message = await client.messages.create(
            model=MODEL, max_tokens=1200, system=PROPHET_SYSTEM,
            messages=[{"role": "user", "content": prompt + profile_note}],
        )
        analyse = "".join(b.text for b in message.content if b.type == "text").strip()
        return {"mode": market["mode"], "analyse": analyse, "market": market["data"]}
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erreur d'analyse : {str(e)}")


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
